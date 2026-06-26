#!/usr/bin/env python3
"""LitSieve local backend.

Serves the LitSieve frontend and implements /api/analyze.
Set DEEPSEEK_API_KEY to enable model-backed analysis.
"""

from __future__ import annotations

import cgi
import io
import json
import mimetypes
import os
import re
import shutil
import subprocess
import tempfile
import urllib.error
import urllib.request
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any

try:
    import pdfplumber
except Exception:  # pragma: no cover
    pdfplumber = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

try:
    import docx
except Exception:  # pragma: no cover
    docx = None

try:
    import fitz
except Exception:  # pragma: no cover
    fitz = None

APP_DIR = Path(__file__).resolve().parent
DEEPSEEK_ENDPOINT = os.environ.get("DEEPSEEK_API_BASE", "https://api.deepseek.com/chat/completions")
DEEPSEEK_MODEL = os.environ.get("DEEPSEEK_MODEL", "deepseek-chat")
MODEL_TEXT_CHUNK_CHARS = int(os.environ.get("MODEL_TEXT_CHUNK_CHARS", "18000"))
MODEL_MAX_DIRECT_CHARS = int(os.environ.get("MODEL_MAX_DIRECT_CHARS", "54000"))


class APIError(Exception):
    def __init__(self, status: int, message: str) -> None:
        super().__init__(message)
        self.status = status
        self.message = message


class LitSieveHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args: Any, **kwargs: Any) -> None:
        super().__init__(*args, directory=str(APP_DIR), **kwargs)

    def do_GET(self) -> None:
        if self.path == "/healthz":
            self.send_json(200, {"ok": True, "engine": "litsieve"})
            return
        super().do_GET()

    def do_HEAD(self) -> None:
        if self.path == "/healthz":
            self.send_response(200)
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.end_headers()
            return
        super().do_HEAD()

    def do_POST(self) -> None:
        if self.path != "/api/analyze":
            self.send_error(404, "Unknown API route")
            return

        try:
            payload = parse_multipart(self)
            response = analyze_payload(payload)
            self.send_json(200, response)
        except APIError as exc:
            self.send_json(exc.status, {"error": exc.message, "results": []})
        except Exception as exc:
            self.send_json(500, {"error": str(exc), "results": []})

    def send_json(self, status: int, data: dict[str, Any]) -> None:
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def parse_multipart(handler: LitSieveHandler) -> dict[str, Any]:
    content_type = handler.headers.get("Content-Type", "")
    if not content_type.startswith("multipart/form-data"):
        raise ValueError("Expected multipart/form-data")

    form = cgi.FieldStorage(
        fp=handler.rfile,
        headers=handler.headers,
        environ={
            "REQUEST_METHOD": "POST",
            "CONTENT_TYPE": content_type,
            "CONTENT_LENGTH": handler.headers.get("Content-Length", "0"),
        },
        keep_blank_values=True,
    )

    request_raw = get_field_value(form, "request") or "{}"
    request_data = json.loads(request_raw)

    metadata: list[dict[str, Any]] = []
    for key in sorted((k for k in form.keys() if k.startswith("metadata_")), key=lambda value: int(value.split("_", 1)[1])):
        metadata.append(json.loads(get_field_value(form, key) or "{}"))

    files = []
    file_fields = []
    if "files" in form:
        file_field = form["files"]
        file_fields = file_field if isinstance(file_field, list) else [file_field]
    for item in file_fields:
        if getattr(item, "filename", None):
            files.append({"filename": item.filename, "content": item.file.read()})

    return {"request": request_data, "metadata": metadata, "files": files}


def get_field_value(form: cgi.FieldStorage, key: str) -> str | None:
    if key not in form:
        return None
    field = form[key]
    if isinstance(field, list):
        field = field[0]
    value = field.value
    if isinstance(value, bytes):
        return value.decode("utf-8", errors="replace")
    return value


def analyze_payload(payload: dict[str, Any]) -> dict[str, Any]:
    request = payload["request"]
    materials = build_materials(payload["metadata"], payload["files"])
    if not materials:
        raise APIError(400, "请先上传文献、图片、PDF，或粘贴一段材料。")

    api_key = os.environ.get("DEEPSEEK_API_KEY", "").strip()
    if not api_key:
        raise APIError(503, "DeepSeek 后端尚未配置 DEEPSEEK_API_KEY，系统不能生成正式文献筛选结论。")

    return {"engineMode": "deepseek", "results": analyze_with_deepseek(request, materials, api_key)}


def build_materials(metadata: list[dict[str, Any]], files: list[dict[str, Any]]) -> list[dict[str, Any]]:
    file_iter = iter(files)
    materials = []

    for meta in metadata:
        file_info = next(file_iter, None) if not meta.get("pastedText") else None
        filename = file_info["filename"] if file_info else meta.get("title", "pasted text")
        content = file_info["content"] if file_info else (meta.get("pastedText") or "").encode("utf-8")
        extracted = extract_text(filename, content)
        materials.append(
            {
                "id": meta.get("id") or filename,
                "title": meta.get("title") or strip_extension(filename),
                "type": meta.get("type") or "mixed",
                "filename": filename,
                "text": extracted["text"],
                "pages": extracted.get("pages", []),
                "extraction": extracted["method"],
            }
        )

    return materials


def extract_text(filename: str, content: bytes) -> dict[str, Any]:
    ext = Path(filename).suffix.lower()
    if ext in {".txt", ".md", ".csv", ".tsv"}:
        return {"method": "text", "text": decode_text(content), "pages": []}
    if ext == ".docx":
        return extract_docx(content)
    if ext == ".pdf":
        return extract_pdf(content)
    if ext in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}:
        return extract_image_text(content, ext)
    if not ext:
        return {"method": "pasted", "text": decode_text(content), "pages": []}
    return {"method": "unknown", "text": decode_text(content), "pages": []}


def decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "gb18030", "shift_jis", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def extract_docx(content: bytes) -> dict[str, Any]:
    if docx is None:
        return {"method": "docx-unavailable", "text": "", "pages": []}
    document = docx.Document(io.BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return {"method": "docx", "text": "\n".join(paragraphs), "pages": []}


def extract_pdf(content: bytes) -> dict[str, Any]:
    pages: list[dict[str, Any]] = []

    if pdfplumber is not None:
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for index, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text() or ""
                    pages.append({"page": index, "text": text})
        except Exception:
            pages = []

    if not any(page["text"].strip() for page in pages) and PdfReader is not None:
        try:
            reader = PdfReader(io.BytesIO(content))
            pages = [{"page": index, "text": page.extract_text() or ""} for index, page in enumerate(reader.pages, start=1)]
        except Exception:
            pages = []

    text = "\n\n".join(f"[p.{page['page']}]\n{page['text']}" for page in pages if page["text"].strip())
    if text.strip():
        return {"method": "pdf-text", "text": text, "pages": pages}

    ocr_result = extract_pdf_ocr(content)
    if ocr_result["text"].strip():
        return ocr_result
    return {"method": ocr_result["method"], "text": "", "pages": pages}


def extract_pdf_ocr(content: bytes) -> dict[str, Any]:
    if fitz is None or not shutil.which("tesseract"):
        return {"method": "pdf-ocr-unavailable", "text": "", "pages": []}

    pages: list[dict[str, Any]] = []
    try:
        document = fitz.open(stream=content, filetype="pdf")
        for index, page in enumerate(document, start=1):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_text = extract_image_text(pixmap.tobytes("png"), ".png")["text"]
            pages.append({"page": index, "text": image_text})
    except Exception:
        return {"method": "pdf-ocr-failed", "text": "", "pages": []}

    text = "\n\n".join(f"[p.{page['page']} OCR]\n{page['text']}" for page in pages if page["text"].strip())
    return {"method": "pdf-ocr", "text": text, "pages": pages}


def extract_image_text(content: bytes, ext: str) -> dict[str, Any]:
    tesseract = shutil.which("tesseract")
    if not tesseract:
        return {
            "method": "image-no-ocr-engine",
            "text": "",
            "pages": [],
        }

    suffix = ext if ext.startswith(".") else f".{ext}"
    with tempfile.NamedTemporaryFile(suffix=suffix) as image_file:
        image_file.write(content)
        image_file.flush()
        completed = subprocess.run(
            [tesseract, image_file.name, "stdout", "-l", os.environ.get("TESSERACT_LANG", "eng+chi_sim+jpn")],
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            check=False,
        )
    return {"method": "tesseract", "text": completed.stdout.strip(), "pages": []}


def analyze_with_deepseek(request: dict[str, Any], materials: list[dict[str, Any]], api_key: str) -> list[dict[str, Any]]:
    results = []
    for index, material in enumerate(materials):
        result = analyze_material_with_deepseek(request, material, api_key)
        results.append(normalize_result(result, [material], 0))
    return results


def analyze_material_with_deepseek(request: dict[str, Any], material: dict[str, Any], api_key: str) -> dict[str, Any]:
    text = material["text"].strip()
    if len(text) <= MODEL_MAX_DIRECT_CHARS:
        return call_deepseek_json(build_model_prompt(request, [material]), api_key).get("results", [{}])[0]

    chunk_results = []
    for chunk_index, chunk in enumerate(chunk_text(text, MODEL_TEXT_CHUNK_CHARS), start=1):
        chunk_material = {**material, "text": chunk, "title": f"{material['title']} - chunk {chunk_index}"}
        chunk_prompt = build_model_prompt(
            {**request, "chunkMode": True},
            [chunk_material],
            extra_rules=[
                "This is one continuous chunk from a longer source. Judge only this chunk, name pages or section markers if visible, and do not infer unavailable parts.",
            ],
        )
        chunk_results.append(call_deepseek_json(chunk_prompt, api_key).get("results", [{}])[0])

    final_prompt = build_final_prompt(request, material, chunk_results)
    return call_deepseek_json(final_prompt, api_key).get("results", [{}])[0]


def call_deepseek_json(prompt: str, api_key: str) -> dict[str, Any]:
    body = {
        "model": DEEPSEEK_MODEL,
        "temperature": 0.1,
        "messages": [
            {
                "role": "system",
                "content": "You are LitSieve, an academic literature screening engine. Return strict JSON only.",
            },
            {"role": "user", "content": prompt},
        ],
        "response_format": {"type": "json_object"},
    }
    req = urllib.request.Request(
        DEEPSEEK_ENDPOINT,
        data=json.dumps(body, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=120) as response:
            raw = response.read().decode("utf-8")
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"DeepSeek API error {exc.code}: {detail}") from exc

    payload = json.loads(raw)
    content = payload["choices"][0]["message"]["content"]
    return parse_model_json(content)


def parse_model_json(content: str) -> dict[str, Any]:
    clean = content.strip()
    if clean.startswith("```"):
        clean = re.sub(r"^```(?:json)?\s*", "", clean)
        clean = re.sub(r"\s*```$", "", clean)
    parsed = json.loads(clean)
    if isinstance(parsed, list):
        return {"results": parsed}
    return parsed


def build_model_prompt(request: dict[str, Any], materials: list[dict[str, Any]], extra_rules: list[str] | None = None) -> str:
    compact_materials = []
    for material in materials:
        text = material["text"].strip()
        compact_materials.append(
            {
                "id": material["id"],
                "title": material["title"],
                "type": material["type"],
                "extraction": material["extraction"],
                "text": text,
                "textLength": len(text),
                "hasReadableText": bool(text),
            }
        )

    return json.dumps(
        {
            "task": "Screen uploaded academic and non-academic materials for relevance to the user's research need.",
            "requirements": request,
            "outputSchema": {
                "results": [
                    {
                        "id": "same id as input",
                        "title": "material title",
                        "status": "keep | maybe | reject",
                        "score": "0-100 number",
                        "tags": ["short tags"],
                        "contentSummary": "summary in requested output language",
                        "usableScope": "full work/chapter/section/passage/page scope in requested output language",
                        "selectionReason": "reason for using or not using in requested output language",
                        "limitations": "citation/source limits in requested output language",
                        "segments": [{"label": "chapter/page/section", "score": 0, "text": "short locator or excerpt"}],
                    }
                ]
            },
            "rules": [
                "Always output in the selected outputLanguage.",
                "Read and evaluate the supplied material content before deciding status.",
                "Never tell the user to wait for DeepSeek, OCR, backend, or future analysis. Return the final screening result from the content you received.",
                "Do not require literal keyword appearance when keywordMode is concept.",
                "Translate concepts across languages internally.",
                "If a whole book is not relevant but one chapter is, status should be maybe and usableScope must name that chapter or section.",
                "Only reject after content analysis shows it cannot support the user's need.",
            ]
            + (extra_rules or []),
            "materials": compact_materials,
        },
        ensure_ascii=False,
    )


def build_final_prompt(request: dict[str, Any], material: dict[str, Any], chunk_results: list[dict[str, Any]]) -> str:
    return json.dumps(
        {
            "task": "Combine chunk-level readings into one final literature screening result for the whole uploaded material.",
            "requirements": request,
            "material": {
                "id": material["id"],
                "title": material["title"],
                "type": material["type"],
                "extraction": material["extraction"],
            },
            "chunkReadings": chunk_results,
            "outputSchema": {
                "results": [
                    {
                        "id": material["id"],
                        "title": material["title"],
                        "status": "keep | maybe | reject",
                        "score": "0-100 number",
                        "tags": ["short tags"],
                        "contentSummary": "whole-material summary in requested output language",
                        "usableScope": "full work/chapter/section/passage/page scope in requested output language",
                        "selectionReason": "reason for using or not using in requested output language",
                        "limitations": "citation/source limits in requested output language",
                        "segments": [{"label": "chapter/page/section", "score": 0, "text": "short locator or excerpt"}],
                    }
                ]
            },
            "rules": [
                "Output only strict JSON.",
                "Always output in the selected outputLanguage.",
                "Never mention waiting for OCR, DeepSeek, backend, or future analysis.",
                "If only part of the material is useful, status must be maybe and usableScope must name that part.",
                "If the material is useful as a whole, status should be keep.",
            ],
        },
        ensure_ascii=False,
    )


def chunk_text(text: str, size: int) -> list[str]:
    return [text[index : index + size] for index in range(0, len(text), size)]


def normalize_result(result: dict[str, Any], materials: list[dict[str, Any]], index: int) -> dict[str, Any]:
    material = materials[index] if index < len(materials) else {}
    status = result.get("status") if result.get("status") in {"keep", "maybe", "reject"} else "maybe"
    return {
        "id": result.get("id") or material.get("id") or f"material-{index + 1}",
        "title": result.get("title") or material.get("title") or f"Material {index + 1}",
        "status": status,
        "score": int(result.get("score") or 60),
        "tags": result.get("tags") or [],
        "contentSummary": result.get("contentSummary") or "",
        "usableScope": result.get("usableScope") or "",
        "selectionReason": result.get("selectionReason") or "",
        "limitations": result.get("limitations") or "",
        "segments": result.get("segments") or [],
    }


def local_result(request: dict[str, Any], material: dict[str, Any]) -> dict[str, Any]:
    text = material["text"].strip()
    terms = expand_terms(request.get("keyword", ""))
    score = relevance_score(text, terms)
    if not text:
        status = "maybe"
        score = max(score, 55)
        summary = "该材料未抽取到可分析正文。"
        scope = "正式部署时由 DeepSeek 与 OCR 服务读取全文后给出章节、页码或片段范围。"
        reason = "未抽取到可分析正文，系统不能生成筛选结论。"
    else:
        status = "keep" if score >= request.get("threshold", 68) else "maybe" if score >= 30 else "reject"
        summary = summarize_text(material["type"], text)
        scope = usable_scope(status)
        reason = local_reason(status, terms, text)

    return {
        "id": material["id"],
        "title": material["title"],
        "status": status,
        "score": score,
        "tags": [term for term in terms[:5] if term],
        "contentSummary": summary,
        "usableScope": scope,
        "selectionReason": reason,
        "limitations": "该结果不能替代 DeepSeek 后端的正式全文分析。",
        "segments": make_segments(text),
    }


def expand_terms(keyword: str) -> list[str]:
    base = [term.strip().lower() for term in re.split(r"[\s,，、;；/]+", keyword) if len(term.strip()) > 1]
    concept_map = {
        "文学理论": ["文学理论", "文艺理论", "文学批评", "literary theory", "structuralism", "结构主义", "feminism", "女性主义", "deconstruction", "解构主义"],
        "literary": ["文学理论", "文艺理论", "literary theory", "literary criticism"],
        "敬语": ["敬语", "敬語", "keigo", "honorific", "politeness", "待遇表現", "尊敬語", "謙譲語", "丁寧語"],
    }
    expanded = list(base)
    for key, terms in concept_map.items():
        if any(key in term for term in base):
            expanded.extend(terms)
    return list(dict.fromkeys(expanded))


def relevance_score(text: str, terms: list[str]) -> int:
    haystack = text.lower()
    hits = sum(1 for term in terms if term and term.lower() in haystack)
    method_hits = sum(1 for term in ["研究", "分析", "考察", "method", "data", "資料", "調査"] if term.lower() in haystack)
    return max(12, min(95, 18 + hits * 18 + method_hits * 6))


def summarize_text(material_type: str, text: str) -> str:
    clean = re.sub(r"\s+", " ", text).strip()
    return f"{type_label(material_type)}材料。可读正文显示：{clean[:220]}{'...' if len(clean) > 220 else ''}"


def usable_scope(status: str) -> str:
    if status == "keep":
        return "可作为核心材料进入精读，优先核对研究对象、方法、结论和引用页码。"
    if status == "maybe":
        return "可能存在局部可用章节或片段，建议限定引用范围后使用。"
    return "当前抽取正文未显示与研究需求的直接支撑关系。"


def local_reason(status: str, terms: list[str], text: str) -> str:
    matched = [term for term in terms if term and term.lower() in text.lower()]
    if status == "reject":
        return "本地抽取文本未发现足够的主题、概念或研究对象线索。"
    if matched:
        return f"本地抽取文本中出现相关概念线索：{'、'.join(matched[:5])}。"
    return "本地抽取文本显示可能相关，但需要 DeepSeek 完整分析确认概念对应关系。"


def make_segments(text: str) -> list[dict[str, Any]]:
    if not text.strip():
        return [{"label": "全文", "score": 55, "text": ""}]
    parts = [part.strip() for part in re.split(r"(?=第[一二三四五六七八九十\d]+[章节節])|\n{2,}", text) if part.strip()]
    return [{"label": f"片段 {index + 1}", "score": 60, "text": part[:160]} for index, part in enumerate(parts[:5])]


def type_label(material_type: str) -> str:
    return {
        "paper": "论文/期刊",
        "book": "专著/编著",
        "news": "新闻/报纸",
        "report": "报告",
        "image": "图片/扫描件",
        "snippet": "文字片段",
        "mixed": "混合材料",
    }.get(material_type, "材料")


def strip_extension(filename: str) -> str:
    return Path(filename).stem.replace("_", " ")


def main() -> None:
    port = int(os.environ.get("PORT", "4173"))
    host = os.environ.get("HOST", "127.0.0.1")
    server = ThreadingHTTPServer((host, port), LitSieveHandler)
    print(f"LitSieve running at http://{host}:{port}")
    server.serve_forever()


if __name__ == "__main__":
    main()
