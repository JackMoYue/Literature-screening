#!/usr/bin/env python3
"""LitSieve local backend.

Serves the LitSieve frontend and implements /api/analyze.
Set DEEPSEEK_API_KEY to enable model-backed analysis.
"""

from __future__ import annotations

import cgi
import io
import json
import mimetypes
import os
import re
import shutil
import subprocess
import tempfile
import urllib.error
import urllib.request
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any

try:
    import pdfplumber
except Exception:  # pragma: no cover
    pdfplumber = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

try:
    import docx
except Exception:  # pragma: no cover
    docx = None

APP_DIR = Path(__file__).resolve().parent
DEEPSEEK_ENDPOINT = os.environ.get("DEEPSEEK_API_BASE", "https://api.deepseek.com/chat/completions")
DEEPSEEK_MODEL = os.environ.get("DEEPSEEK_MODEL", "deepseek-chat")


class LitSieveHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args: Any, **kwargs: Any) -> None:
        super().__init__(*args, directory=str(APP_DIR), **kwargs)

    def do_GET(self) -> None:
        if self.path == "/healthz":
            self.send_json(200, {"ok": True, "engine": "litsieve"})
            return
        super().do_GET()

    def do_HEAD(self) -> None:
        if self.path == "/healthz":
            self.send_response(200)
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.end_headers()
            return
        super().do_HEAD()

    def do_POST(self) -> None:
        if self.path != "/api/analyze":
            self.send_error(404, "Unknown API route")
            return

        try:
            payload = parse_multipart(self)
            response = analyze_payload(payload)
            self.send_json(200, response)
        except Exception as exc:  # Keep frontend usable during early prototyping.
            self.send_json(500, {"error": str(exc), "results": []})

    def send_json(self, status: int, data: dict[str, Any]) -> None:
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def parse_multipart(handler: LitSieveHandler) -> dict[str, Any]:
    content_type = handler.headers.get("Content-Type", "")
    if not content_type.startswith("multipart/form-data"):
        raise ValueError("Expected multipart/form-data")

    form = cgi.FieldStorage(
        fp=handler.rfile,
        headers=handler.headers,
        environ={
            "REQUEST_METHOD": "POST",
            "CONTENT_TYPE": content_type,
            "CONTENT_LENGTH": handler.headers.get("Content-Length", "0"),
        },
        keep_blank_values=True,
    )

    request_raw = get_field_value(form, "request") or "{}"
    request_data = json.loads(request_raw)

    metadata: list[dict[str, Any]] = []
    for key in sorted((k for k in form.keys() if k.startswith("metadata_")), key=lambda value: int(value.split("_", 1)[1])):
        metadata.append(json.loads(get_field_value(form, key) or "{}"))

    files = []
    file_fields = []
    if "files" in form:
        file_field = form["files"]
        file_fields = file_field if isinstance(file_field, list) else [file_field]
    for item in file_fields:
        if getattr(item, "filename", None):
            files.append({"filename": item.filename, "content": item.file.read()})

    return {"request": request_data, "metadata": metadata, "files": files}


def get_field_value(form: cgi.FieldStorage, key: str) -> str | None:
    if key not in form:
        return None
    field = form[key]
    if isinstance(field, list):
        field = field[0]
    value = field.value
    if isinstance(value, bytes):
        return value.decode("utf-8", errors="replace")
    return value


def analyze_payload(payload: dict[str, Any]) -> dict[str, Any]:
    request = payload["request"]
    materials = build_materials(payload["metadata"], payload["files"])

    api_key = os.environ.get("DEEPSEEK_API_KEY", "").strip()
    if api_key:
        return {"results": analyze_with_deepseek(request, materials, api_key)}

    return {
        "engineMode": "local",
        "warning": "DEEPSEEK_API_KEY is not set; returned local extraction-based analysis.",
        "results": [local_result(request, material) for material in materials],
    }


def build_materials(metadata: list[dict[str, Any]], files: list[dict[str, Any]]) -> list[dict[str, Any]]:
    file_iter = iter(files)
    materials = []

    for meta in metadata:
        file_info = next(file_iter, None) if not meta.get("pastedText") else None
        filename = file_info["filename"] if file_info else meta.get("title", "pasted text")
        content = file_info["content"] if file_info else (meta.get("pastedText") or "").encode("utf-8")
        extracted = extract_text(filename, content)
        materials.append(
            {
                "id": meta.get("id") or filename,
                "title": meta.get("title") or strip_extension(filename),
                "type": meta.get("type") or "mixed",
                "filename": filename,
                "text": extracted["text"],
                "pages": extracted.get("pages", []),
                "extraction": extracted["method"],
            }
        )

    return materials


def extract_text(filename: str, content: bytes) -> dict[str, Any]:
    ext = Path(filename).suffix.lower()
    if ext in {".txt", ".md", ".csv", ".tsv"}:
        return {"method": "text", "text": decode_text(content), "pages": []}
    if ext == ".docx":
        return extract_docx(content)
    if ext == ".pdf":
        return extract_pdf(content)
    if ext in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}:
        return extract_image_text(content, ext)
    if not ext:
        return {"method": "pasted", "text": decode_text(content), "pages": []}
    return {"method": "unknown", "text": decode_text(content), "pages": []}


def decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "gb18030", "shift_jis", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def extract_docx(content: bytes) -> dict[str, Any]:
    if docx is None:
        return {"method": "docx-unavailable", "text": "", "pages": []}
    document = docx.Document(io.BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return {"method": "docx", "text": "\n".join(paragraphs), "pages": []}


def extract_pdf(content: bytes) -> dict[str, Any]:
    pages: list[dict[str, Any]] = []

    if pdfplumber is not None:
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for index, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text() or ""
                    pages.append({"page": index, "text": text})
        except Exception:
            pages = []

    if not any(page["text"].strip() for page in pages) and PdfReader is not None:
        try:
            reader = PdfReader(io.BytesIO(content))
            pages = [{"page": index, "text": page.extract_text() or ""} for index, page in enumerate(reader.pages, start=1)]
        except Exception:
            pages = []

    text = "\n\n".join(f"[p.{page['page']}]\n{page['text']}" for page in pages if page["text"].strip())
    return {"method": "pdf", "text": text, "pages": pages}


def extract_image_text(content: bytes, ext: str) -> dict[str, Any]:
    tesseract = shutil.which("tesseract")
    if not tesseract:
        return {
            "method": "image-no-ocr-engine",
            "text": "",
            "pages": [],
        }

    suffix = ext if ext.startswith(".") else f".{ext}"
    with tempfile.NamedTemporaryFile(suffix=suffix) as image_file:
        image_file.write(content)
        image_file.flush()
        completed = subprocess.run(
            [tesseract, image_file.name, "stdout", "-l", os.environ.get("TESSERACT_LANG", "eng+chi_sim+jpn")],
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            check=False,
        )
    return {"method": "tesseract", "text": completed.stdout.strip(), "pages": []}


def analyze_with_deepseek(request: dict[str, Any], materials: list[dict[str, Any]], api_key: str) -> list[dict[str, Any]]:
    prompt = build_model_prompt(request, materials)
    body = {
        "model": DEEPSEEK_MODEL,
        "temperature": 0.1,
        "messages": [
            {
                "role": "system",
                "content": "You are LitSieve, an academic literature screening engine. Return strict JSON only.",
            },
            {"role": "user", "content": prompt},
        ],
        "response_format": {"type": "json_object"},
    }
    req = urllib.request.Request(
        DEEPSEEK_ENDPOINT,
        data=json.dumps(body, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=120) as response:
            raw = response.read().decode("utf-8")
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"DeepSeek API error {exc.code}: {detail}") from exc

    payload = json.loads(raw)
    content = payload["choices"][0]["message"]["content"]
    parsed = json.loads(content)
    results = parsed.get("results", parsed if isinstance(parsed, list) else [])
    return [normalize_result(result, materials, index) for index, result in enumerate(results)]


def build_model_prompt(request: dict[str, Any], materials: list[dict[str, Any]]) -> str:
    compact_materials = []
    for material in materials:
        text = material["text"].strip()
        compact_materials.append(
            {
                "id": material["id"],
                "title": material["title"],
                "type": material["type"],
                "extraction": material["extraction"],
                "text": text[:24000],
            }
        )

    return json.dumps(
        {
            "task": "Screen uploaded academic and non-academic materials for relevance to the user's research need.",
            "requirements": request,
            "outputSchema": {
                "results": [
                    {
                        "id": "same id as input",
                        "title": "material title",
                        "status": "keep | maybe | reject",
                        "score": "0-100 number",
                        "tags": ["short tags"],
                        "contentSummary": "summary in requested output language",
                        "usableScope": "full work/chapter/section/passage/page scope in requested output language",
                        "selectionReason": "reason for using or not using in requested output language",
                        "limitations": "citation/source limits in requested output language",
                        "segments": [{"label": "chapter/page/section", "score": 0, "text": "short locator or excerpt"}],
                    }
                ]
            },
            "rules": [
                "Always output in the selected outputLanguage.",
                "Do not require literal keyword appearance when keywordMode is concept.",
                "Translate concepts across languages internally.",
                "If a whole book is not relevant but one chapter is, status should be maybe and usableScope must name that chapter or section.",
                "Only reject after content analysis shows it cannot support the user's need.",
            ],
            "materials": compact_materials,
        },
        ensure_ascii=False,
    )


def normalize_result(result: dict[str, Any], materials: list[dict[str, Any]], index: int) -> dict[str, Any]:
    material = materials[index] if index < len(materials) else {}
    status = result.get("status") if result.get("status") in {"keep", "maybe", "reject"} else "maybe"
    return {
        "id": result.get("id") or material.get("id") or f"material-{index + 1}",
        "title": result.get("title") or material.get("title") or f"Material {index + 1}",
        "status": status,
        "score": int(result.get("score") or 60),
        "tags": result.get("tags") or [],
        "contentSummary": result.get("contentSummary") or "",
        "usableScope": result.get("usableScope") or "",
        "selectionReason": result.get("selectionReason") or "",
        "limitations": result.get("limitations") or "",
        "segments": result.get("segments") or [],
    }


def local_result(request: dict[str, Any], material: dict[str, Any]) -> dict[str, Any]:
    text = material["text"].strip()
    terms = expand_terms(request.get("keyword", ""))
    score = relevance_score(text, terms)
    if not text:
        status = "maybe"
        score = max(score, 55)
        summary = "该材料已经进入全文读取流程。当前本地后端没有可用的图片文字识别引擎，因此只返回开发环境占位结果。"
        scope = "正式部署时由 DeepSeek 与 OCR 服务读取全文后给出章节、页码或片段范围。"
        reason = "未抽取到可分析正文，本地后端不作排除结论。"
    else:
        status = "keep" if score >= request.get("threshold", 68) else "maybe" if score >= 30 else "reject"
        summary = summarize_text(material["type"], text)
        scope = usable_scope(status)
        reason = local_reason(status, terms, text)

    return {
        "id": material["id"],
        "title": material["title"],
        "status": status,
        "score": score,
        "tags": [term for term in terms[:5] if term],
        "contentSummary": summary,
        "usableScope": scope,
        "selectionReason": reason,
        "limitations": "本地后端结果用于开发测试；设置 DEEPSEEK_API_KEY 后会返回模型生成的正式分析。",
        "segments": make_segments(text),
    }


def expand_terms(keyword: str) -> list[str]:
    base = [term.strip().lower() for term in re.split(r"[\s,，、;；/]+", keyword) if len(term.strip()) > 1]
    concept_map = {
        "文学理论": ["文学理论", "文艺理论", "文学批评", "literary theory", "structuralism", "结构主义", "feminism", "女性主义", "deconstruction", "解构主义"],
        "literary": ["文学理论", "文艺理论", "literary theory", "literary criticism"],
        "敬语": ["敬语", "敬語", "keigo", "honorific", "politeness", "待遇表現", "尊敬語", "謙譲語", "丁寧語"],
    }
    expanded = list(base)
    for key, terms in concept_map.items():
        if any(key in term for term in base):
            expanded.extend(terms)
    return list(dict.fromkeys(expanded))


def relevance_score(text: str, terms: list[str]) -> int:
    haystack = text.lower()
    hits = sum(1 for term in terms if term and term.lower() in haystack)
    method_hits = sum(1 for term in ["研究", "分析", "考察", "method", "data", "資料", "調査"] if term.lower() in haystack)
    return max(12, min(95, 18 + hits * 18 + method_hits * 6))


def summarize_text(material_type: str, text: str) -> str:
    clean = re.sub(r"\s+", " ", text).strip()
    return f"{type_label(material_type)}材料。可读正文显示：{clean[:220]}{'...' if len(clean) > 220 else ''}"


def usable_scope(status: str) -> str:
    if status == "keep":
        return "可作为核心材料进入精读，优先核对研究对象、方法、结论和引用页码。"
    if status == "maybe":
        return "可能存在局部可用章节或片段，建议限定引用范围后使用。"
    return "当前抽取正文未显示与研究需求的直接支撑关系。"


def local_reason(status: str, terms: list[str], text: str) -> str:
    matched = [term for term in terms if term and term.lower() in text.lower()]
    if status == "reject":
        return "本地抽取文本未发现足够的主题、概念或研究对象线索。"
    if matched:
        return f"本地抽取文本中出现相关概念线索：{'、'.join(matched[:5])}。"
    return "本地抽取文本显示可能相关，但需要 DeepSeek 完整分析确认概念对应关系。"


def make_segments(text: str) -> list[dict[str, Any]]:
    if not text.strip():
        return [{"label": "全文", "score": 55, "text": ""}]
    parts = [part.strip() for part in re.split(r"(?=第[一二三四五六七八九十\d]+[章节節])|\n{2,}", text) if part.strip()]
    return [{"label": f"片段 {index + 1}", "score": 60, "text": part[:160]} for index, part in enumerate(parts[:5])]


def type_label(material_type: str) -> str:
    return {
        "paper": "论文/期刊",
        "book": "专著/编著",
        "news": "新闻/报纸",
        "report": "报告",
        "image": "图片/扫描件",
        "snippet": "文字片段",
        "mixed": "混合材料",
    }.get(material_type, "材料")


def strip_extension(filename: str) -> str:
    return Path(filename).stem.replace("_", " ")


def main() -> None:
    port = int(os.environ.get("PORT", "4173"))
    host = os.environ.get("HOST", "127.0.0.1")
    server = ThreadingHTTPServer((host, port), LitSieveHandler)
    print(f"LitSieve running at http://{host}:{port}")
    server.serve_forever()


if __name__ == "__main__":
    main()
